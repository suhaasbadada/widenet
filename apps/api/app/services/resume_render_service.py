import logging
import os
import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt


_logger = logging.getLogger(__name__)

_REQUIRED_FIELDS = {
    "name",
    "contact_number",
    "links",
    "summary",
    "skills",
    "experience",
    "projects",
    "education",
}


class ResumeRenderValidationError(ValueError):
    """Raised when the render request payload is invalid."""


class ResumeRenderFailedError(RuntimeError):
    """Raised when DOCX rendering or PDF conversion fails."""


@dataclass
class RenderedDocxArtifact:
    docx_path: str
    temp_dir: tempfile.TemporaryDirectory[str]


@dataclass
class RenderedPdfArtifact:
    pdf_path: str
    temp_dir: tempfile.TemporaryDirectory[str]


class _CallableList(list[Any]):
    """List that is also callable to support template patterns like values and values()."""

    def __call__(self) -> "_CallableList":
        return self


class _TemplateLinks(_CallableList):
    """Template-friendly links container supporting both list and dot-access patterns."""

    def __init__(
        self,
        items: list[str],
        email: str = "",
        linkedin: str = "",
        github: str = "",
        portfolio: str = "",
        website: str = "",
    ) -> None:
        super().__init__(items)
        self.email = _TemplateLinkEntry(email)
        self.linkedin = _TemplateLinkEntry(linkedin)
        self.github = _TemplateLinkEntry(github)
        self.portfolio = _TemplateLinkEntry(portfolio)
        self.website = _TemplateLinkEntry(website)


class _TemplateLinkEntry:
    """Expose link values with the `show` attribute expected by template fields."""

    def __init__(self, value: str = "") -> None:
        normalized = (value or "").strip()
        self.show = normalized
        self.href = normalized
        self.url = normalized

    def __str__(self) -> str:
        return self.show


class _TemplateMapping:
    """Mapping wrapper that is template-friendly for dot-access and dict-like helpers."""

    def __init__(self, data: dict[str, Any]) -> None:
        self._data = data

    def __getattribute__(self, name: str) -> Any:
        # Prefer payload keys over helper attributes so template fields like
        # `skill.items` resolve to the actual resume value, not mapping helpers.
        if name not in {"_data", "__class__", "__dict__", "__weakref__"}:
            data = object.__getattribute__(self, "_data")
            if isinstance(data, dict) and name in data:
                return _to_template_context(data[name])
        return object.__getattribute__(self, name)

    def __getitem__(self, key: str) -> Any:
        return _to_template_context(self._data[key])

    def __iter__(self):
        for key in self._data:
            yield key

    def __len__(self) -> int:
        return len(self._data)

    @property
    def keys(self) -> _CallableList:
        return _CallableList([_to_template_context(v) for v in self._data.keys()])

    @property
    def values(self) -> _CallableList:
        return _CallableList([_to_template_context(v) for v in self._data.values()])

    @property
    def items(self) -> _CallableList:
        return _CallableList([(k, _to_template_context(v)) for k, v in self._data.items()])

    def get(self, key: str, default: Any = None) -> Any:
        if key in self._data:
            return _to_template_context(self._data[key])
        return default

    def __getattr__(self, name: str) -> Any:
        if name in self._data:
            return _to_template_context(self._data[name])
        raise AttributeError(name)


def render_resume_to_docx(
    resume_json: Any,
    template_path: str,
    file_name: str = "resume.docx",
) -> RenderedDocxArtifact:
    """Render a resume JSON into DOCX using a template."""
    validated_resume = _normalize_resume_for_template(_validate_resume_json(resume_json))
    validated_template_path = _validate_template_path(template_path)
    _normalize_file_name(file_name)
    temp_dir, docx_output = _render_docx_artifact(validated_resume, validated_template_path)

    return RenderedDocxArtifact(docx_path=str(docx_output), temp_dir=temp_dir)


def render_resume_to_pdf(
    resume_json: Any,
    template_path: str,
    file_name: str = "resume.pdf",
) -> RenderedPdfArtifact:
    """Render a resume JSON into PDF using DOCX template + LibreOffice conversion."""
    validated_resume = _normalize_resume_for_template(_validate_resume_json(resume_json))
    validated_template_path = _validate_template_path(template_path)
    _normalize_pdf_file_name(file_name)

    temp_dir, docx_output = _render_docx_artifact(validated_resume, validated_template_path)
    work_dir = Path(temp_dir.name)
    pdf_output = docx_output.with_suffix(".pdf")
    libreoffice_profile_dir = work_dir / "libreoffice-profile"
    libreoffice_profile_dir.mkdir(parents=True, exist_ok=True)

    # Template loops can leave trailing empty paragraphs that sometimes force
    # one dangling line onto page 2 after DOCX->PDF conversion.
    _trim_trailing_empty_paragraphs(docx_output)
    _normalize_section_leading_spacing(docx_output)

    soffice_bin = _resolve_libreoffice_binary()
    command = [
        soffice_bin,
        f"-env:UserInstallation=file://{libreoffice_profile_dir}",
        "--headless",
        "--convert-to",
        "pdf:writer_pdf_Export",
        "--outdir",
        str(work_dir),
        str(docx_output),
    ]

    try:
        result = subprocess.run(command, check=False, capture_output=True, text=True)
    except Exception as exc:
        temp_dir.cleanup()
        _logger.exception("Failed to execute LibreOffice command using '%s'.", soffice_bin)
        raise ResumeRenderFailedError("Failed to run LibreOffice for PDF conversion.") from exc

    if result.returncode != 0:
        temp_dir.cleanup()
        _logger.error(
            "LibreOffice PDF conversion failed (exit=%s). stderr=%s",
            result.returncode,
            (result.stderr or "").strip(),
        )
        raise ResumeRenderFailedError("Failed to convert DOCX to PDF.")

    if not pdf_output.exists():
        temp_dir.cleanup()
        _logger.error("LibreOffice reported success but PDF was not created at '%s'.", pdf_output)
        raise ResumeRenderFailedError("PDF output file was not created.")

    return RenderedPdfArtifact(pdf_path=str(pdf_output), temp_dir=temp_dir)


def _render_docx_artifact(
    validated_resume: dict[str, Any],
    validated_template_path: Path,
) -> tuple[tempfile.TemporaryDirectory[str], Path]:
    """Render validated resume data into a temporary DOCX artifact."""

    temp_dir = tempfile.TemporaryDirectory(prefix="widenet-resume-render-")
    work_dir = Path(temp_dir.name)
    docx_output = work_dir / "rendered_resume.docx"

    template_context = _to_template_context(validated_resume)

    try:
        template = DocxTemplate(str(validated_template_path))
        template.render(template_context)
        template.save(str(docx_output))
    except Exception as exc:
        temp_dir.cleanup()
        _logger.exception("Failed to render DOCX from template '%s'.", validated_template_path)
        raise ResumeRenderFailedError("Failed to render DOCX from template.") from exc

    if not docx_output.exists():
        temp_dir.cleanup()
        _logger.error("DOCX output file was not created at '%s'.", docx_output)
        raise ResumeRenderFailedError("DOCX output file was not created.")

    return temp_dir, docx_output


def _resolve_libreoffice_binary() -> str:
    """Resolve a usable LibreOffice CLI binary path."""
    env_bin = os.environ.get("LIBREOFFICE_BIN", "").strip()
    if env_bin:
        candidate = Path(env_bin).expanduser()
        if candidate.is_file():
            return str(candidate)
        found = shutil.which(env_bin)
        if found:
            return found

    for name in ("soffice", "libreoffice"):
        found = shutil.which(name)
        if found:
            return found

    common_paths = [
        Path("/Applications/LibreOffice.app/Contents/MacOS/soffice"),
        Path("/usr/bin/soffice"),
        Path("/usr/local/bin/soffice"),
        Path("/opt/homebrew/bin/soffice"),
        Path("/snap/bin/libreoffice"),
    ]
    for candidate in common_paths:
        if candidate.is_file():
            return str(candidate)

    raise ResumeRenderFailedError(
        "LibreOffice CLI was not found. Install LibreOffice or set LIBREOFFICE_BIN."
    )


def _validate_resume_json(resume_json: Any) -> dict[str, Any]:
    if not isinstance(resume_json, dict):
        raise ResumeRenderValidationError("Field 'resume_json' must be a JSON object.")

    missing = sorted(_REQUIRED_FIELDS - set(resume_json.keys()))
    if missing:
        raise ResumeRenderValidationError(
            f"Field 'resume_json' is missing required keys: {', '.join(missing)}."
        )

    return resume_json


def _validate_template_path(template_path: str) -> Path:
    raw_path = Path(template_path).expanduser()
    project_root = Path(__file__).resolve().parents[2]

    # Resolve relative template paths from project root so runtime cwd does not matter.
    candidates: list[Path] = []
    if raw_path.is_absolute():
        candidates.append(raw_path)
    else:
        candidates.append(Path.cwd() / raw_path)
        candidates.append(project_root / raw_path)

        # Common project convention for templates under app/resume-templates.
        candidates.append(project_root / "app" / raw_path)

        # If caller passed only a filename, also search known template directories.
        if raw_path.parent == Path("."):
            candidates.append(project_root / "resume-templates" / raw_path.name)
            candidates.append(project_root / "app" / "resume-templates" / raw_path.name)

    resolved_path: Path | None = None
    for candidate in candidates:
        candidate_resolved = candidate.resolve()
        if candidate_resolved.is_file():
            resolved_path = candidate_resolved
            break

    if resolved_path is None:
        checked = ", ".join(str(c.resolve()) for c in candidates)
        _logger.error("Template not found. requested='%s' checked=[%s]", template_path, checked)
        raise ResumeRenderValidationError(
            f"Template file does not exist: '{template_path}'."
        )

    path = resolved_path
    if path.suffix.lower() != ".docx":
        raise ResumeRenderValidationError("Field 'template_path' must point to a .docx file.")
    return path


def _normalize_file_name(file_name: str | None) -> str:
    name = (file_name or "resume.docx").strip() or "resume.docx"
    if not name.lower().endswith(".docx"):
        name = f"{name}.docx"
    return name


def _normalize_pdf_file_name(file_name: str | None) -> str:
    name = (file_name or "resume.pdf").strip() or "resume.pdf"
    if not name.lower().endswith(".pdf"):
        name = f"{name}.pdf"
    return name


def _trim_trailing_empty_paragraphs(docx_path: Path) -> None:
    """Remove trailing blank paragraphs to reduce PDF pagination drift."""
    try:
        document = Document(str(docx_path))
    except Exception:
        # If the file cannot be loaded, keep conversion behavior unchanged.
        return

    removed = 0
    while document.paragraphs:
        last = document.paragraphs[-1]
        if last.text.strip():
            break
        parent = last._element.getparent()
        if parent is None:
            break
        parent.remove(last._element)
        removed += 1

    if removed:
        document.save(str(docx_path))


def _normalize_section_leading_spacing(docx_path: Path) -> None:
    """Reduce spacing drift after section headers for DOCX->PDF conversion."""
    try:
        document = Document(str(docx_path))
    except Exception:
        return

    changed = False
    paragraphs = document.paragraphs
    section_titles = {
        "summary",
        "skills",
        "experience",
        "projects",
        "education",
        "certifications",
    }

    for idx, paragraph in enumerate(paragraphs):
        title = paragraph.text.strip().lower().rstrip(":")
        if title not in section_titles:
            continue

        # Remove blank paragraph runs directly after a section heading.
        probe = idx + 1
        while probe < len(paragraphs) and not paragraphs[probe].text.strip():
            blank_para = paragraphs[probe]
            parent = blank_para._element.getparent()
            if parent is None:
                break
            parent.remove(blank_para._element)
            changed = True
            probe += 1

        # Ensure first visible content line starts without extra leading space.
        if probe < len(paragraphs) and paragraphs[probe].text.strip():
            fmt = paragraphs[probe].paragraph_format
            if fmt.space_before is None or float(fmt.space_before.pt) > 0:
                fmt.space_before = Pt(0)
                changed = True

    if changed:
        document.save(str(docx_path))


def _to_template_context(value: Any) -> Any:
    if isinstance(value, _TemplateLinks):
        return value
    if isinstance(value, _TemplateLinkEntry):
        return value
    if isinstance(value, dict):
        return _TemplateMapping(value)
    if isinstance(value, list):
        return [_to_template_context(item) for item in value]
    return value


def _normalize_resume_for_template(resume: dict[str, Any]) -> dict[str, Any]:
    normalized = dict(resume)
    normalized["skills"] = _normalize_skills_structure(resume.get("skills"))
    normalized["links"] = _normalize_links_structure(
        links=resume.get("links"),
        fallback_email=str(resume.get("email", "")).strip(),
    )
    normalized["experience"] = _normalize_experience_for_template(resume.get("experience"))
    normalized["projects"] = _normalize_projects_for_template(resume.get("projects"))
    normalized["education"] = _normalize_education_for_template(resume.get("education"))
    return normalized


def _to_string(value: Any) -> str:
    return str(value).strip() if value is not None else ""


def _to_string_list(value: Any) -> list[str]:
    if isinstance(value, list):
        return [_to_string(item) for item in value if _to_string(item)]
    text = _to_string(value)
    return [text] if text else []


def _parse_duration_range(duration: str) -> tuple[str, str]:
    text = duration.strip()
    if not text:
        return "", ""

    for separator in (" – ", " - ", "–", "-"):
        if separator in text:
            start, end = text.split(separator, 1)
            return start.strip(), end.strip()
    return text, ""


def _normalize_experience_for_template(experience: Any) -> list[dict[str, Any]]:
    if not isinstance(experience, list):
        return []

    normalized: list[dict[str, Any]] = []
    for item in experience:
        if not isinstance(item, dict):
            continue

        title = _to_string(item.get("title") or item.get("role") or item.get("position"))
        company = _to_string(item.get("company") or item.get("employer"))
        location = _to_string(item.get("location"))

        from_val = _to_string(item.get("from") or item.get("start"))
        to_val = _to_string(item.get("to") or item.get("end"))
        if not from_val and not to_val:
            parsed_start, parsed_end = _parse_duration_range(_to_string(item.get("duration")))
            from_val = parsed_start
            to_val = parsed_end

        bullets = _to_string_list(item.get("bullets"))
        if not bullets:
            bullets = _to_string_list(item.get("points"))
        if not bullets:
            bullets = _to_string_list(item.get("responsibilities"))
        if not bullets:
            bullets = _to_string_list(item.get("description"))

        normalized.append(
            {
                "title": title,
                "company": company,
                "location": location,
                "from": from_val,
                "to": to_val,
                "bullets": bullets,
            }
        )

    return normalized


def _normalize_projects_for_template(projects: Any) -> list[dict[str, Any]]:
    if not isinstance(projects, list):
        return []

    normalized: list[dict[str, Any]] = []
    for item in projects:
        if not isinstance(item, dict):
            continue

        name = _to_string(item.get("name") or item.get("title"))
        subtext = _to_string(item.get("subtext") or item.get("description") or item.get("summary"))

        tech_value = item.get("tech") or item.get("technologies") or item.get("stack")
        if isinstance(tech_value, list):
            tech = ", ".join(_to_string(v) for v in tech_value if _to_string(v))
        else:
            tech = _to_string(tech_value)

        bullets = _to_string_list(item.get("bullets"))
        if not bullets:
            bullets = _to_string_list(item.get("points"))
        if not bullets:
            bullets = _to_string_list(item.get("impact"))
        if not bullets and subtext:
            bullets = [subtext]

        normalized.append(
            {
                "name": name,
                "subtext": subtext,
                "tech": tech,
                "bullets": bullets,
            }
        )

    return normalized


def _normalize_education_for_template(education: Any) -> list[dict[str, Any]]:
    if not isinstance(education, list):
        return []

    normalized: list[dict[str, Any]] = []
    for item in education:
        if not isinstance(item, dict):
            continue

        school = _to_string(item.get("school") or item.get("institution") or item.get("university"))
        location = _to_string(item.get("location"))
        degree = _to_string(item.get("degree"))
        major = _to_string(item.get("major") or item.get("field") or item.get("specialization"))
        gpa = _to_string(item.get("gpa") or item.get("cgpa"))
        from_value = _to_string(item.get("from") or item.get("start"))
        to_value = _to_string(item.get("to") or item.get("end"))

        year = _to_string(item.get("year"))
        if year and not to_value:
            to_value = year

        normalized.append(
            {
                "school": school,
                "location": location,
                "degree": degree,
                "major": major,
                "gpa": gpa,
                "from": from_value,
                "to": to_value,
            }
        )

    return normalized


def _normalize_links_structure(links: Any, fallback_email: str = "") -> _TemplateLinks:
    items: list[str] = []
    email = ""
    linkedin = ""
    github = ""
    portfolio = ""
    website = ""

    if isinstance(links, dict):
        email = str(links.get("email", "")).strip()
        linkedin = str(links.get("linkedin", "")).strip()
        github = str(links.get("github", "")).strip()
        portfolio = str(links.get("portfolio", "")).strip()
        website = str(links.get("website", "")).strip()

        raw_items = links.get("items")
        if isinstance(raw_items, list):
            items.extend(str(item).strip() for item in raw_items if str(item).strip())

        for value in links.values():
            if isinstance(value, str) and value.strip():
                items.append(value.strip())
    elif isinstance(links, list):
        items.extend(str(item).strip() for item in links if str(item).strip())
    elif isinstance(links, str) and links.strip():
        items.append(links.strip())

    deduped: list[str] = []
    seen: set[str] = set()
    for item in items:
        key = item.lower()
        if key not in seen:
            seen.add(key)
            deduped.append(item)

        normalized = item.lower()
        if not email:
            if normalized.startswith("mailto:") and len(item) > len("mailto:"):
                email = item[len("mailto:") :]
            elif "@" in item and "://" not in item and "/" not in item:
                email = item

        if "linkedin.com" in normalized and not linkedin:
            linkedin = item
        if "github.com" in normalized and not github:
            github = item
        if any(host in normalized for host in ("behance.net", "dribbble.com")) and not portfolio:
            portfolio = item
        if normalized.startswith("http") and not website:
            website = item

    if not email and fallback_email:
        email = fallback_email

    return _TemplateLinks(
        items=deduped,
        email=email,
        linkedin=linkedin,
        github=github,
        portfolio=portfolio,
        website=website,
    )


def _normalize_skills_structure(skills: Any) -> list[dict[str, Any]]:
    if isinstance(skills, list):
        normalized_list: list[dict[str, Any]] = []
        for entry in skills:
            if isinstance(entry, dict):
                category = str(entry.get("category", "")).strip()
                items_raw = entry.get("items", [])
                items = [str(item).strip() for item in (items_raw or []) if str(item).strip()]
                normalized_list.append({"category": category, "items": items})
                continue

            # Handle accidental tuple-pair shape like:
            # [('category', 'Languages'), ('items', ['Python'])]
            if isinstance(entry, (list, tuple)):
                try:
                    entry_dict = dict(entry)
                except (TypeError, ValueError):
                    continue
                category = str(entry_dict.get("category", "")).strip()
                items_raw = entry_dict.get("items", [])
                items = [str(item).strip() for item in (items_raw or []) if str(item).strip()]
                normalized_list.append({"category": category, "items": items})
        return normalized_list

    if isinstance(skills, dict):
        normalized_list = []
        for category, items_raw in skills.items():
            items = [str(item).strip() for item in (items_raw or []) if str(item).strip()]
            normalized_list.append({"category": str(category), "items": items})
        return normalized_list

    return []
